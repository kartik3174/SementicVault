import os
import fitz  # PyMuPDF
from docx import Document as DocxDocument

def load_txt(file_path: str) -> str:
    """Load plain text files (TXT, MD)."""
    with open(file_path, "r", encoding="utf-8", errors="ignore") as f:
        return f.read()

def load_markdown(file_path: str) -> str:
    """Load Markdown files."""
    return load_txt(file_path)

def load_docx(file_path: str) -> str:
    """Load MS Word DOCX files."""
    doc = DocxDocument(file_path)
    full_text = []
    for para in doc.paragraphs:
        full_text.append(para.text)
    
    # Also extract text from tables
    for table in doc.tables:
        for row in table.rows:
            row_text = [cell.text.strip() for cell in row.cells if cell.text.strip()]
            if row_text:
                full_text.append(" | ".join(row_text))
                
    return "\n".join(full_text)

def load_pdf(file_path: str) -> str:
    """Load PDF files using PyMuPDF and preserve page markers."""
    doc = fitz.open(file_path)
    full_text = []
    for i, page in enumerate(doc):
        text = page.get_text()
        full_text.append(f"[Page {i + 1}]\n{text}")
    return "\n\n".join(full_text)

def load_document(file_path: str, mime_type: str = None) -> str:
    """
    Unified loader helper.
    Dispatches to the correct loader based on file extension or mime type.
    """
    ext = os.path.splitext(file_path)[1].lower()
    
    if ext in [".txt", ".text"]:
        return load_txt(file_path)
    elif ext in [".md", ".markdown"]:
        return load_markdown(file_path)
    elif ext == ".docx":
        return load_docx(file_path)
    elif ext == ".pdf":
        return load_pdf(file_path)
    else:
        # Fallback to plain text load
        try:
            return load_txt(file_path)
        except Exception as e:
            raise ValueError(f"Unsupported file format '{ext}' or loading failed: {str(e)}")
