"""
Loader module for Word (.docx) documents using python-docx.
"""
import logging
from typing import Tuple
from docx import Document as DocxDocument

logger = logging.getLogger("SemanticVault.Loader.Docx")

class DocxLoader:
    """Loader for Word (.docx) files that extracts paragraphs and tabular text."""
    
    def __init__(self, file_path: str):
        self.file_path = file_path

    def load(self) -> Tuple[str, int]:
        """
        Extracts content from a Microsoft Word .docx file.
        
        Note:
            DOCX files do not contain explicit hard page-break tags unless manually inserted, 
            as paging is done dynamically by the rendering word processor. Therefore, we treat 
            the document as a single dynamic flow (Page Count = 1).
            
        Returns:
            Tuple[str, int]: A tuple containing the extracted text and the page count (always 1).
        """
        logger.info(f"Extracting text from DOCX: {self.file_path}")
        try:
            doc = DocxDocument(self.file_path)
            full_text = []
            
            # 1. Extract paragraphs
            for para in doc.paragraphs:
                if para.text:
                    full_text.append(para.text)
                    
            # 2. Extract tables text to ensure high retrieval capability for structured data
            for t_idx, table in enumerate(doc.tables):
                logger.info(f"Parsing table #{t_idx+1} in DOCX file")
                for row in table.rows:
                    row_text = [cell.text.strip() for cell in row.cells if cell.text.strip()]
                    if row_text:
                        # Standardize tabular row cells with pipe separators
                        full_text.append(" | ".join(row_text))
            
            combined_text = "\n".join(full_text)
            logger.info(f"Successfully loaded DOCX document. Characters: {len(combined_text)}")
            return combined_text, 1
            
        except Exception as e:
            logger.error(f"Failed to load DOCX document from {self.file_path}: {e}", exc_info=True)
            raise RuntimeError(f"python-docx failed to process document: {str(e)}")
